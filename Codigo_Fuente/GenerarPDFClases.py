from docx import Document
import sys
import os
import comtypes.client
import re
import docx
import pkg_resources
pkg_resources.require("xlrd==1.2.0")
import xlrd
import xlsxwriter

class FormattedDocument:
    #Defining constructor for class
    def __init__(self, path, ExcDoc=None):
        self.path = path
        self.plainText = self.setPlainText()
        self.formatAmount = self.setFormatAmount()
        self.formatSet = self.setFormatSet()
        self.formatDictionary = {}
        self.docxDoc = docx.Document(self.path)
        self.ExcDoc = ExcDoc

    #Defining methods to get or modify class attributes
    def getPath(self):
        return self.path

    def setPath(self,stringPath):
        self.path = stringPath
    
    def getExcDoc(self):
        return self.ExcDoc
    
    def getDocxDoc(self):
        return self.docxDoc
    
    def setPlainText(self):
        #Creating docx document
        doc = docx.Document(self.path)
        #Saving each paragraph in the doxc document
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        #Returning data
        return '\n'.join(fullText)
    
    def getPlainText(self):
        return self.plainText
    
    def setFormatAmount(self):
        #Getting the plain text information from the document
        toStrDoc = self.getPlainText()
        #Searching all format patterns using regex
        allFormatsList = re.findall("(\[\w*\d*\])",toStrDoc)
        #Creating a list with all formats with no repetition
        FormatsNoRep = []
        for element in allFormatsList:
            if (element not in FormatsNoRep):
                FormatsNoRep.append(element)
        #Returning data
        return len(FormatsNoRep)
    
    def getFormatAmount(self):
        return self.formatAmount
    
    def setFormatSet(self):
        #Getting the plain text information from the document
        toStrDoc = self.getPlainText()
        #Searching all format patterns using regex
        allFormatsList = re.findall("(\[\w*\d*\])",toStrDoc)
        #Creating a list with all formats with no repetition
        FormatsNoRep = []
        for element in allFormatsList:
            if (element not in FormatsNoRep):
                FormatsNoRep.append(element)
        #Returning data
        return FormatsNoRep
    
    def getFormatSet(self):
        return self.formatSet
    
    def DictFromExc(self, pathToExc):
        # Creating Excel document objects
        loc = (pathToExc)
        wb = xlrd.open_workbook(loc)
        sheet = wb.sheet_by_index(0)
        sheet.cell_value(0, 0)
        #Getting row number
        rows = sheet.nrows
        #Creating the main dictionary using the Excel file information
        dictionary = {}
        rowsList = []
        for x in range(6,rows):
            rowsList.append(sheet.row_values(x))
        for row in rowsList:
            for x in range(len(row)):
                if isinstance(row[x], float):
                    row[x] = str(int(row[x]))
        for row in rowsList:
            dictionary['\\' +row[0][:-1] + '\\]'] = row[1:]
        #Returning dict
        self.formatDictionary = dictionary
    
    def lgthNamePdfTemp(self,pathToExc):
        # Creating document
        loc = (pathToExc)
        wb = xlrd.open_workbook(loc)
        sheet = wb.sheet_by_index(0)
        sheet.cell_value(0, 0)
        #Getting row number
        rows = sheet.nrows
        #Creating a dictionary that will contain some information about the user's preferences
        #Creating an auxiliary list
        NameAndPdf = {}
        rowsList = []
        #Adding URL information
        NameAndPdf[sheet.row_values(0)[0]] = sheet.row_values(0)[1]
        #Adding 'NOMBRE ARCHIVO', 'CREAR PDF', 'NOMBRE PDF' with its values to the dictionary using the aux list
        for x in range(3,7):
            rowsList.append(sheet.row_values(x))
        for row in rowsList:
            NameAndPdf[row[0]] = row[1:]
        
        #Returning data
        return len(rowsList[0]), NameAndPdf
        
    
    def getFormatDict(self):
        return self.formatDictionary
    
    def docx_replace_regex(self, doc_obj, regex , replace):
        #Checking all paragraphs
        for p in doc_obj.paragraphs:
            #Searching for the pattern using regex
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text
        #Searching for patterns in all the tables in the document
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.docx_replace_regex(cell, regex , replace)
        
    
    def create_pdf(self,docPath,name):
        #This function creates and saves a PDF from a docx file
        wdFormatPDF = 17
        in_file = os.path.abspath(docPath)
        out_file = os.path.abspath(name)
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_file)
        doc.SaveAs(out_file, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()


class User:
    #Defining constructor
    def __init__(self,document,profiles=None):
        self.document = document
        self.profiles = profiles

    def changeDocument(self,dirPath):
        #Initializing information
        pathToExc = self.document.getExcDoc()
        self.document.DictFromExc(pathToExc)
        dictionary = self.document.getFormatDict()

        #Replacement process
        (length,dictNamePdf) = self.document.lgthNamePdfTemp(pathToExc)
        for x in range(length-1):
            self.document.docxDoc = docx.Document(self.document.path)
            for word, replacement in dictionary.items():
                doc = self.document.getDocxDoc()
                word_re=re.compile(word)
                self.document.docx_replace_regex(doc,word_re , replacement[x])
            
            #Saving and creating pdf if the user desires it in the Excel file contents
            fullPath = dirPath +dictNamePdf["NOMBRE ARCHIVO"][x]+ ".docx"
            doc.save(fullPath)
            if dictNamePdf["CREAR PDF"][x] == "SI":
                pdfpath = dirPath +dictNamePdf["NOMBRE PDF"][x]+".pdf"
                self.document.create_pdf(fullPath,pdfpath)

def getTemplateDict(pathToExc):
        # Creating document
        loc = (pathToExc)
        wb = xlrd.open_workbook(loc)
        sheet = wb.sheet_by_index(0)
        sheet.cell_value(0, 0)
        #Getting row number
        rows = sheet.nrows
        tDict = {}
        #Add template
        tDict[sheet.row_values(0)[0]] = sheet.row_values(0)[1]
        
        return tDict


def writeProfile(templatePath,templateName,dirPath, profileName):
    #Create document
    document = FormattedDocument(templatePath)
    formatList = list(document.getFormatSet())
    #Creating and setting up Excel file
    newExc = dirPath+profileName+".xlsx"
    workbook = xlsxwriter.Workbook(newExc)
    worksheet = workbook.add_worksheet()
    worksheet.set_column(0, 0, 25)
    worksheet.set_column(0, 1, 25)
    bold = workbook.add_format({'bold': True})
    cell_format = workbook.add_format()
    cell_format.set_align('fill')
    #Writting main information to the Excel document with indications to the user
    worksheet.write('A1', 'URL',bold)
    worksheet.write('B1', templatePath,cell_format)
    worksheet.write('A2', "PLANTILLA",bold)
    worksheet.write('B2', templateName)
    worksheet.write('B3', "VALUE_1",bold)
    worksheet.write('A3', 'KEY',bold)
    worksheet.write('A4', 'NOMBRE ARCHIVO')
    worksheet.write('B4', 'Ingrese un nombre')
    worksheet.write('A5', 'CREAR PDF')
    worksheet.write('B5', 'SI/NO')
    worksheet.write('A6', 'NOMBRE PDF')
    worksheet.write('B6', 'Ingrese un nombre')
    index = 0
    for x in range(7,len(formatList)+7):
        cell = 'A'+str(x)
        sideCell = 'B' + str(x)
        worksheet.write(cell, formatList[index])
        worksheet.write(sideCell, "Ingrese un valor")
        index += 1
    #Closing...
    workbook.close()

def generateDocuments(profilePath, dirPath):
    #Creating main dictionary
    docDict = getTemplateDict(profilePath)
    #Getting URL information to find the desired docx template
    dPath = docDict["URL"]
    #Creating FormattedDocument object
    Documento = FormattedDocument(dPath,profilePath)
    #Creating User object
    Usuario = User(Documento)
    #Generating formatted document with the user's indications according to the profile
    Usuario.changeDocument(dirPath)


def openProfile(profilePath):
    #Opening the profile in 'profilePath'
    os.startfile(profilePath)

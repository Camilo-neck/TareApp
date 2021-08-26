from docx import Document
import sys
import os
import comtypes.client
import re
import docx
import pkg_resources
pkg_resources.require("xlrd==1.2.0")
import xlrd

from datetime import datetime as dt
import locale

class FormattedDocument:
    #Defining constructor for class
    def __init__(self, path, exc_doc=None):
        self.path = path
        self.plain_text = self.set_plain_text()
        self.format_amount = self.set_format_amount()
        self.format_set = self.set_format_set()
        self.format_dictionary = {}
        self.docx_doc = docx.Document(self.path)
        self.exc_doc =exc_doc 
        #Stablish region for the language used in date formating
        locale.setlocale(locale.LC_TIME, '')
        #Create actual_date object
        self.actual_date = dt.now()

    #Defining methods to get or modify class attributes
    def get_path(self):
        return self.path

    def set_path(self,stringPath):
        self.path = stringPath
    
    def get_exc_doc(self):
        return self.exc_doc
    
    def get_docx_doc(self):
        return self.docx_doc
    
    def set_plain_text(self):
        #Creating docx document
        doc = docx.Document(self.path)
        #Saving each paragraph in the doxc document
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        #Returning data
        return '\n'.join(full_text)
    
    def getplain_text(self):
        return self.plain_text
    
    def set_format_amount(self):
        #Getting the plain text information from the document
        to_str_doc = self.getplain_text()
        #Searching all format patterns using regex
        all_formats_list = re.findall("(\[\w*\d*\])",to_str_doc)
        #Creating a list with all formats with no repetition
        formats_no_rep = []
        for element in all_formats_list:
            if (element not in formats_no_rep):
                formats_no_rep.append(element)
        #Returning data
        return len(formats_no_rep)
    
    def get_format_amount(self):
        return self.format_amount
    
    def set_format_set(self):
        #Getting the plain text information from the document
        to_str_doc = self.getplain_text()
        #Searching all format patterns using regex
        all_formats_list = re.findall("(\[\w*\d*\])",to_str_doc)
        #Creating a list with all formats with no repetition
        formats_no_rep = []
        for element in all_formats_list:
            if (element not in formats_no_rep):
                formats_no_rep.append(element)
        #Returning data
        return formats_no_rep
    
    def get_format_set(self):
        return self.format_set
    
    def dict_from_exc(self, path_to_exc):
        # Creating Excel document objects
        loc = (path_to_exc)
        wb = xlrd.open_workbook(loc)
        sheet = wb.sheet_by_index(0)
        sheet.cell_value(0, 0)
        #Getting row number
        rows = sheet.nrows
        #Creating the main dictionary using the Excel file information
        dictionary = {}
        rows_list = []
        for x in range(7,rows):
            rows_list.append(sheet.row_values(x))
        for row in rows_list:
            for x in range(len(row)):
                if isinstance(row[x], float):
                    row[x] = str(int(row[x]))
        for row in rows_list:
            #set dict keys and values (if date format is presented in value it will be changued via mapping)   
            dictionary['\\' +row[0][:-1] + '\\]'] = list(map(lambda v: self.actual_date.strftime(v),row[1:]))

        #Returning dict
        self.format_dictionary = dictionary
        
    
    def lgth_name_pdf_temp(self,path_to_exc):
        # Creating document
        loc = (path_to_exc)
        wb = xlrd.open_workbook(loc)
        sheet = wb.sheet_by_index(0)
        sheet.cell_value(0, 0)
        #Creating a dictionary that will contain some information about the user's preferences
        #Creating an auxiliary list
        name_and_pdf = {}
        rows_list = []
        #Adding URL information
        name_and_pdf[sheet.row_values(0)[0]] = sheet.row_values(0)[1]
        #Adding 'NOMBRE ARCHIVO', 'CREAR PDF', 'NOMBRE PDF' with its values to the dictionary using the aux list
        for x in range(4,8):
            rows_list.append(sheet.row_values(x))
        for row in rows_list:
            name_and_pdf[row[0]] = list(map(lambda v: self.actual_date.strftime(v),row[1:]))
        
        #Returning data
        return len(rows_list[0]), name_and_pdf
        
    
    def get_format_dict(self):
        return self.format_dictionary
    
    def docx_replace_regex(self, doc_obj, regex , replace):
        #Checking all paragraphs
        for p in doc_obj.paragraphs:
            #Searching for the pattern using regex
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text
        #Searching for patterns in all the tables in the document
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.docx_replace_regex(cell, regex , replace)
        
    
    def create_pdf(self,doc_path,name):
        #This function creates and saves a PDF from a docx file
        wd_format_pdf = 17
        in_file = os.path.abspath(doc_path)
        out_file = os.path.abspath(name)
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_file)
        doc.SaveAs(out_file, FileFormat=wd_format_pdf)
        doc.Close()
        word.Quit()
    
    def validate_template(self):
        main_dictionary = self.get_format_dict()
        format_list = self.getformat_set().copy()
        for x in range(len(format_list)):
            new_str = "\\" + format_list[x][:-1] + "\\]"
            format_list[x] = new_str
        if set(main_dictionary) == set(format_list):
            return True
        else:
            return False


class User:
    #Defining constructor
    def __init__(self,document,profiles=None):
        self.document = document
        self.profiles = profiles
        self.status = 0
        self.log = ""

    def change_document(self,dir_path):
        #Initializing information
        path_to_exc = self.document.getexc_doc()
        self.document.dict_from_exc(path_to_exc)
        dictionary = self.document.get_format_dict()
        try:
            if self.document.validate_template() and os.path.isdir(dir_path):
                #Replacement process
                (length,dict_name_pdf) = self.document.lgth_name_pdf_temp(path_to_exc)
                for x in range(length-1):
                    self.document.docx_doc = docx.Document(self.document.path)
                    for word, replacement in dictionary.items():
                        doc = self.document.get_docx_doc()
                        word_re=re.compile(word)
                        self.document.docx_replace_regex(doc,word_re , replacement[x])
                    
                    #Saving and creating pdf if the user desires it in the Excel file contents
                    full_path = dir_path + "/"+dict_name_pdf["NOMBRE ARCHIVO"][x]+ ".docx"
                    doc.save(full_path)
                    if dict_name_pdf["CREAR PDF"][x] == "SI":
                        
                        pdfpath = dir_path + "/"+dict_name_pdf["NOMBRE PDF"][x]+".pdf"
                        self.document.create_pdf(full_path,pdfpath)
                self.log="Documentos generados correctamente"
            else:
                self.status = 1
                self.log="Hubo un problema al generar los documentos,\npor favor, revise la información de su perfil"
        except Exception as e:
            self.status = 1
            self.log = str(e)
